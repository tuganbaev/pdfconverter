"""
Utility functions for PDF conversion with Cyrillic support
"""
import io
import os
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.colors import black
from docx import Document as DocxDocument


def register_cyrillic_fonts():
    """Register fonts that support Cyrillic characters"""
    font_configs = [
        {
            'name': 'DejaVuSans',
            'paths': [
                '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
                '/usr/share/fonts/TTF/DejaVuSans.ttf',
            ]
        },
        {
            'name': 'DejaVuSans-Bold',
            'paths': [
                '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
                '/usr/share/fonts/TTF/DejaVuSans-Bold.ttf',
            ]
        },
        {
            'name': 'Ubuntu',
            'paths': [
                '/usr/share/fonts/truetype/ubuntu/Ubuntu-R.ttf',
            ]
        }
    ]
    
    registered_fonts = []
    
    for config in font_configs:
        for path in config['paths']:
            if os.path.exists(path):
                try:
                    pdfmetrics.registerFont(TTFont(config['name'], path))
                    registered_fonts.append(config['name'])
                    print(f"Registered font: {config['name']} from {path}")
                    break
                except Exception as e:
                    print(f"Failed to register {config['name']}: {e}")
                    continue
    
    return registered_fonts


def create_cyrillic_styles():
    """Create paragraph styles that support Cyrillic text"""
    # Try to register fonts
    registered_fonts = register_cyrillic_fonts()
    
    # Choose the best available font
    if 'DejaVuSans' in registered_fonts:
        base_font = 'DejaVuSans'
        bold_font = 'DejaVuSans-Bold' if 'DejaVuSans-Bold' in registered_fonts else 'DejaVuSans'
    elif 'Ubuntu' in registered_fonts:
        base_font = 'Ubuntu'
        bold_font = 'Ubuntu'
    else:
        # Fallback to built-in fonts (limited Cyrillic support)
        base_font = 'Helvetica'
        bold_font = 'Helvetica-Bold'
    
    print(f"Using fonts: base={base_font}, bold={bold_font}")
    
    # Create custom styles
    styles = {
        'normal': ParagraphStyle(
            'CyrillicNormal',
            fontName=base_font,
            fontSize=12,
            leading=16,
            spaceAfter=12,
            textColor=black,
            leftIndent=0,
            rightIndent=0,
        ),
        'heading': ParagraphStyle(
            'CyrillicHeading',
            fontName=bold_font,
            fontSize=16,
            leading=20,
            spaceAfter=16,
            spaceBefore=16,
            textColor=black,
            leftIndent=0,
            rightIndent=0,
        ),
        'title': ParagraphStyle(
            'CyrillicTitle',
            fontName=bold_font,
            fontSize=20,
            leading=24,
            spaceAfter=20,
            spaceBefore=10,
            textColor=black,
            leftIndent=0,
            rightIndent=0,
            alignment=1,  # Center
        )
    }
    
    return styles


def convert_docx_to_pdf_with_cyrillic(docx_path):
    """Convert DOCX to PDF with proper Cyrillic support"""
    try:
        # Read DOCX
        doc = DocxDocument(docx_path)
        
        # Create PDF buffer
        buffer = io.BytesIO()
        pdf_doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            topMargin=72,
            bottomMargin=72,
            leftMargin=72,
            rightMargin=72
        )
        
        # Create styles with Cyrillic support
        styles = create_cyrillic_styles()
        
        # Build content
        story = []
        
        # Add document title
        story.append(Paragraph("Converted Document", styles['title']))
        story.append(Spacer(1, 20))
        
        # Process paragraphs
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            try:
                # Escape XML/HTML characters
                text = escape_xml_chars(text)
                
                # Determine style based on paragraph properties
                if is_heading_paragraph(para, text):
                    story.append(Paragraph(text, styles['heading']))
                else:
                    story.append(Paragraph(text, styles['normal']))
                
                # Add some spacing between paragraphs
                if i < len(doc.paragraphs) - 1:
                    story.append(Spacer(1, 6))
                    
            except Exception as para_error:
                print(f"Error processing paragraph: {para_error}")
                # Add fallback content
                try:
                    safe_text = f"[Paragraph {i+1}: Content processing error]"
                    story.append(Paragraph(safe_text, styles['normal']))
                    story.append(Spacer(1, 6))
                except:
                    continue
        
        # If no content was processed
        if len(story) <= 2:
            story.append(Paragraph("Document processed but no readable content found.", styles['normal']))
        
        # Build PDF
        pdf_doc.build(story)
        
        # Get PDF content
        pdf_content = buffer.getvalue()
        buffer.close()
        
        return pdf_content
        
    except Exception as e:
        print(f"Cyrillic conversion error: {e}")
        return None


def escape_xml_chars(text):
    """Escape characters that might cause XML parsing issues"""
    replacements = {
        '&': '&amp;',
        '<': '&lt;',
        '>': '&gt;',
        '"': '&quot;',
        "'": '&apos;',
    }
    
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    
    return text


def is_heading_paragraph(para, text):
    """Determine if a paragraph should be treated as a heading"""
    # Check paragraph style
    if para.style and para.style.name.startswith('Heading'):
        return True
    
    # Check text characteristics
    if len(text) < 100 and (text.isupper() or text.istitle()):
        return True
    
    # Check if it's a short line that might be a heading
    if len(text.split()) <= 8 and len(text) < 80:
        # Look for common heading patterns
        heading_indicators = ['глава', 'раздел', 'часть', 'введение', 'заключение', 'содержание']
        if any(indicator in text.lower() for indicator in heading_indicators):
            return True
    
    return False


def transliterate_cyrillic(text):
    """Transliterate Cyrillic text to Latin as fallback"""
    replacements = {
        # Lowercase
        'а': 'a', 'б': 'b', 'в': 'v', 'г': 'g', 'д': 'd', 'е': 'e', 'ё': 'yo',
        'ж': 'zh', 'з': 'z', 'и': 'i', 'й': 'y', 'к': 'k', 'л': 'l', 'м': 'm',
        'н': 'n', 'о': 'o', 'п': 'p', 'р': 'r', 'с': 's', 'т': 't', 'у': 'u',
        'ф': 'f', 'х': 'h', 'ц': 'ts', 'ч': 'ch', 'ш': 'sh', 'щ': 'sch',
        'ъ': '', 'ы': 'y', 'ь': '', 'э': 'e', 'ю': 'yu', 'я': 'ya',
        
        # Uppercase  
        'А': 'A', 'Б': 'B', 'В': 'V', 'Г': 'G', 'Д': 'D', 'Е': 'E', 'Ё': 'YO',
        'Ж': 'ZH', 'З': 'Z', 'И': 'I', 'Й': 'Y', 'К': 'K', 'Л': 'L', 'М': 'M',
        'Н': 'N', 'О': 'O', 'П': 'P', 'Р': 'R', 'С': 'S', 'Т': 'T', 'У': 'U',
        'Ф': 'F', 'Х': 'H', 'Ц': 'TS', 'Ч': 'CH', 'Ш': 'SH', 'Щ': 'SCH',
        'Ъ': '', 'Ы': 'Y', 'Ь': '', 'Э': 'E', 'Ю': 'YU', 'Я': 'YA'
    }
    
    result = text
    for cyrillic, latin in replacements.items():
        result = result.replace(cyrillic, latin)
    
    return result